import os
import pandas as pd
from docx import Document
import warnings

# Suppress warnings
warnings.filterwarnings("ignore")

def read_docx(file_path):
    print(f"--- Reading {file_path} ---")
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
        
        # Also read tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                print(" | ".join(row_text))
    except Exception as e:
        print(f"Error reading docx: {e}")
    print("\n")

def read_xls(file_path):
    print(f"--- Reading {file_path} ---")
    try:
        # Try reading as excel
        # Check extension
        if file_path.endswith('.xls'):
            df = pd.read_excel(file_path, engine='xlrd')
        else:
            df = pd.read_excel(file_path)
        
        print(df.to_string())
    except Exception as e:
        print(f"Error reading xls: {e}")
    print("\n")

def main():
    files = [f for f in os.listdir('.') if os.path.isfile(f)]
    
    for f in files:
        if f.endswith('.docx'):
            read_docx(f)
        elif f.endswith('.xls') or f.endswith('.xlsx'):
            read_xls(f)

if __name__ == "__main__":
    main()
